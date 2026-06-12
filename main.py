import csv
import re
import sys

import requests
from docx import Document
from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtGui import QAction, QColor
from PyQt6.QtWidgets import (
    QApplication,
    QComboBox,
    QFileDialog,
    QHeaderView,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QTableWidget,
    QTableWidgetItem,
    QToolBar,
    QVBoxLayout,
    QWidget,
)


# ФОНОВЫЙ ПОТОК ДЛЯ ПЕРЕВОДА
# ФОНОВЫЙ ПОТОК ДЛЯ ПЕРЕВОДА
class QAThread(QThread):
    progress_signal = pyqtSignal(int, str)

    def __init__(self, data, glossary, server, model):
        super().__init__()
        self.data = data
        self.glossary = glossary
        self.server = server
        self.model = model

    def run(self):
        # Цикл должен обхватывать ВСЮ логику генерации и отправки
        for i, (orig, trans) in enumerate(self.data):
            if not orig.strip() or not trans.strip():
                continue

            glossary_str = ", ".join(
                [f"'{k}': '{v}'" for k, v in self.glossary.items()]
            )
            prompt = (
                f"Оригинал: {orig}. Перевод: {trans}. "
                f"Глоссарий: {{{glossary_str}}}. Найди, есть ли в оригинале термины из глоссария. "
                "Если в оригинале НЕТ ни одного термина из глоссария — ответь ТОЛЬКО словом 'SKIP'. "
                "Если термины есть, проверь правильность их перевода с учетом падежей. Если всё верно — ответь 'OK', если ошибка — 'ERROR'."
            )

            # Выбор сервера
            if self.server == "Ollama":
                url = "http://127.0.0.1:11434/api/generate"
                payload = {"model": self.model, "prompt": prompt, "stream": False}
            elif self.server == "LM Studio":
                url = "http://127.0.0.1:1234/v1/chat/completions"
                payload = {
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                }
            else:
                self.progress_signal.emit(i, "ERROR")
                continue

            try:
                response = requests.post(url, json=payload)
                if response.status_code == 200:
                    result = (
                        response.json().get("response", "")
                        if self.server == "Ollama"
                        else response.json()["choices"][0]["message"]["content"]
                    )
                    self.progress_signal.emit(i, result.strip())
                else:
                    self.progress_signal.emit(i, "ERROR")
            except Exception:
                self.progress_signal.emit(i, "ERROR")


class TranslationThread(QThread):
    progress_signal = pyqtSignal(int, str)
    finished_signal = pyqtSignal()

    def __init__(self, texts, target_lang, server, model):
        super().__init__()
        self.texts = texts
        self.target_lang = target_lang
        self.server = server
        self.model = model

    def run(self):
        # Аналогично, цикл обхватывает всё
        for i, text in enumerate(self.texts):
            if not text.strip():
                continue

            prompt = f"Переведи этот текст на {self.target_lang}. Выведи ТОЛЬКО перевод. Никаких объяснений, кавычек или markdown-разметки:\n\n{text}"

            if self.server == "Ollama":
                url = "http://127.0.0.1:11434/api/generate"
                payload = {"model": self.model, "prompt": prompt, "stream": False}
            elif self.server == "LM Studio":
                url = "http://127.0.0.1:1234/v1/chat/completions"
                payload = {
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                }
            else:
                self.progress_signal.emit(i, "ОШИБКА: Неизвестный сервер")
                continue

            try:
                response = requests.post(url, json=payload)
                if response.status_code == 200:
                    result = (
                        response.json().get("response", "")
                        if self.server == "Ollama"
                        else response.json()["choices"][0]["message"]["content"]
                    )
                    self.progress_signal.emit(i, result.strip())
                else:
                    self.progress_signal.emit(
                        i, f"ОШИБКА: Сбой API ({response.status_code})"
                    )
            except Exception as e:
                self.progress_signal.emit(i, f"ОШИБКА: {str(e)}")

        self.finished_signal.emit()


# ГЛАВНОЕ ОКНО
class TranslatorApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Desktop Translator")
        self.resize(1000, 700)
        self.glossary = {}

        # Инструменты
        toolbar = QToolBar("Main Toolbar")
        self.addToolBar(toolbar)

        btn_load_doc = QAction("Загрузить документ", self)
        btn_load_mem = QAction("Загрузить память (CSV)", self)
        btn_translate = QAction("Перевести", self)
        btn_qa = QAction("QA Проверка", self)
        btn_export = QAction("Экспорт", self)

        # Привязка кнопок к функциям
        btn_load_doc.triggered.connect(self.load_document)
        btn_translate.triggered.connect(self.start_translation)
        btn_qa.triggered.connect(self.start_qa)
        btn_load_mem.triggered.connect(self.load_glossary)
        btn_export.triggered.connect(self.export_document)

        toolbar.addAction(btn_load_doc)
        toolbar.addAction(btn_load_mem)
        toolbar.addAction(btn_translate)
        toolbar.addAction(btn_qa)
        toolbar.addAction(btn_export)

        self.server_combo = QComboBox()
        self.server_combo.addItems(["Ollama", "LM Studio"])
        toolbar.addWidget(self.server_combo)

        self.model_input = QLineEdit()
        self.model_input.setPlaceholderText("Название модели")
        toolbar.addWidget(self.model_input)

        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["Английский", "Русский", "Испанский", "Немецкий"])
        toolbar.addWidget(self.lang_combo)

        # Таблица
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        self.table = QTableWidget(0, 2)
        header = self.table.horizontalHeader()
        if header:
            header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
            header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)

        self.table.setWordWrap(True)
        layout.addWidget(self.table)

    def load_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл",
            "",
            "Документы Word (*.docx);;Текстовые файлы (*.txt)",
        )

        if file_path:
            try:
                paragraphs = []
                # Развилка форматов:
                if file_path.endswith(".docx"):
                    doc = Document(file_path)
                    self.current_file_path = (
                        file_path  # Сохраняю путь от оригинального файла
                    )
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                elif file_path.endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as file:
                        paragraphs = [
                            line.strip() for line in file.readlines() if line.strip()
                        ]

                # Нарезал текст и вставил в таблицу
                sentences = self.segment_into_sentences(paragraphs)

                self.table.setRowCount(len(sentences))
                for i, sentence in enumerate(sentences):
                    item = QTableWidgetItem(sentence)
                    self.table.setItem(i, 0, item)
                    self.table.resizeRowToContents(i)

            except Exception as e:
                QMessageBox.critical(
                    self, "Ошибка", f"Не удалось загрузить документ:\n{str(e)}"
                )

    def segment_into_sentences(self, paragraphs):
        sentences = []
        for paragraph in paragraphs:
            splits = re.split(r"(?<=[.!?]) +", paragraph)
            for s in splits:
                if s.strip():
                    sentences.append(s.strip())
        return sentences

    def start_translation(self):
        # Если таблица пустая, ничего не делать
        if self.table.rowCount() == 0:
            QMessageBox.warning(self, "Внимание", "Сначала загрузите документ!")
            return

        target_lang = self.lang_combo.currentText()
        server = self.server_combo.currentText()
        model = self.model_input.text().strip()

        # Собрал все тексты из левой колонки
        texts = []
        for i in range(self.table.rowCount()):
            item = self.table.item(i, 0)
            texts.append(item.text() if item else "")

        # Поменял заголовок окна
        self.setWindowTitle("AI Desktop Translator - ПЕРЕВОД В ПРОЦЕССЕ...")

        # Запускаю фоновый поток
        self.translation_thread = TranslationThread(texts, target_lang, server, model)
        self.translation_thread.progress_signal.connect(self.update_translation_cell)
        self.translation_thread.finished_signal.connect(self.translation_finished)
        self.translation_thread.start()

    def update_translation_cell(self, row, text):
        # Записываются ответы ИИ в правую колонку
        item = QTableWidgetItem(text)
        self.table.setItem(row, 1, item)
        self.table.resizeRowToContents(row)  # Подгоняем высоту строки

    def translation_finished(self):
        self.setWindowTitle("AI Translate")
        QMessageBox.information(self, "Готово", "Перевод успешно завершен!")

    def export_document(self):
        if not hasattr(self, "current_file_path"):
            QMessageBox.warning(self, "Внимание", "Сначала загрузите документ!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить файл",
            "",
            "Документы Word (*.docx);;Текстовые файлы (*.txt)",
        )

        if not file_path:
            return

        try:
            # Развилка форматов:
            if file_path.endswith(".docx"):
                original_doc = Document(self.current_file_path)
                for paragraph, row in zip(
                    original_doc.paragraphs, range(self.table.rowCount())
                ):
                    item = self.table.item(row, 1)  # Текст из правой колонки
                    if item:
                        paragraph.text = item.text()

                original_doc.save(file_path)

            elif file_path.endswith(".txt"):
                with open(file_path, "w", encoding="utf-8") as file:
                    for i in range(self.table.rowCount()):
                        item = self.table.item(i, 1)
                        if item:
                            file.write(item.text() + "\n")

            QMessageBox.information(self, "Готово", "Экспорт успешно завершен!")
        except Exception as e:
            QMessageBox.critical(
                self, "Ошибка", f"Не удалось сохранить документ:\n{str(e)}"
            )

    def load_glossary(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл CSV",
            "",
            "Файлы CSV (*.csv)",
        )
        if file_path:
            try:
                glossary = {}
                with open(file_path, newline="", encoding="utf-8") as csvfile:
                    reader = csv.reader(csvfile)
                    for row in reader:
                        if len(row) == 2:  # Убедился, что каждый ряд содержит 2 столбца
                            original_term, translation = row
                            glossary[original_term] = translation

                # Вывел сообщение об успешном окончании загрузки
                QMessageBox.information(
                    self,
                    "Готово",
                    f"Загруженный глоссарий содержит {len(glossary)} терминов.",
                )
                self.glossary = glossary

            except Exception as e:
                QMessageBox.critical(
                    self, "Ошибка", f"Не удалось загрузить глоссарий:\n{str(e)}"
                )

    def start_qa(self):
        # Если таблица пуста, ничего не произойдёт
        if self.table.rowCount() == 0:
            QMessageBox.warning(self, "Внимание", "Сначала загрузите документ!")
            return

        server = self.server_combo.currentText()
        model = self.model_input.text().strip()

        # Собирал все данные из таблицы
        data = []
        for i in range(self.table.rowCount()):
            original_item = self.table.item(i, 0)
            translation_item = self.table.item(i, 1)
            if original_item and translation_item:
                data.append((original_item.text(), translation_item.text()))

        # запускаю фоновый поток
        self.qa_thread = QAThread(data, self.glossary, server, model)
        self.qa_thread.progress_signal.connect(self.update_qa_cell)
        self.qa_thread.start()

    def update_qa_cell(self, row, status):
        item = self.table.item(row, 1)  # Ячейка перевода
        if item:
            if status == "ERROR":
                item.setBackground(QColor("#FFCCCC"))  # Красный фон для ошибок
            elif status == "OK":
                item.setBackground(
                    QColor("#CCFFCC")
                )  # Зеленый фон для успешных проверок
            else:
                # Пропуск действия для остальных случаев
                pass


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TranslatorApp()
    window.show()
    sys.exit(app.exec())
